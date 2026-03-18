# -*- coding: utf-8 -*-
"""
公文生成器

这个文件是公文生成的核心模块，负责创建Word文档并设置格式
"""

# 导入系统模块，用于处理系统相关的功能
import sys
# 导入io模块，用于处理输入输出编码
import io
# 导入datetime模块，用于处理日期时间
from datetime import datetime
# 导入类型提示模块，让代码更清晰
from typing import Dict, List, Optional
# 导入正则表达式模块，用于标题识别
import re

# 设置标准输出的编码为UTF-8，防止中文乱码
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# 导入python-docx库的Document类，用于创建Word文档
from docx import Document
# 导入python-docx库的单位类，用于设置字体大小、页边距等
from docx.shared import Pt, Cm, Twips, RGBColor
# 导入文本对齐和行间距的枚举类型
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
# 导入表格对齐的枚举类型
from docx.enum.table import WD_TABLE_ALIGNMENT
# 导入XML命名空间相关的工具，用于设置中文字体
from docx.oxml.ns import qn, nsmap
# 导入XML元素创建工具
from docx.oxml import OxmlElement


class OfficialDocumentGenerator:
    """
    党政机关公文生成器类
    
    这是整个项目的核心类，负责创建Word文档并按照国家标准设置格式
    """
    
    # ========== 字体定义 ==========
    # 定义宋体字体名称，用于页码
    FONT_SONGTI = '宋体'
    # 定义仿宋字体名称，用于正文、发文字号
    FONT_FANGSONG = '仿宋'
    # 定义黑体字体名称，用于一级标题、密级
    FONT_HEITI = '黑体'
    # 定义楷体字体名称，用于二级标题、签发人姓名
    FONT_KAITI = '楷体'
    # 定义方正小标宋_GBK字体名称，用于发文机关标志、标题
    FONT_XIAOBIAOSONG = '方正小标宋_GBK'
    
    # 初始化公文生成器
    def __init__(self):
        """
        初始化公文生成器
        
        当创建这个类的实例时，会自动执行以下操作：
        1. 创建一个新的Word文档
        2. 设置页面格式（纸张大小、页边距等）
        3. 设置文档默认样式
        """
        # 创建一个新的Word文档对象
        self.doc = Document()
        # 调用方法设置页面格式
        self._setup_page()
        # 调用方法设置文档样式
        self._setup_styles()
    
    # 设置页面格式（私有方法）
    def _setup_page(self):
        """
        设置页面格式（私有方法）
        
        按照公司格式要求设置：
        - 纸张大小：A4 (210mm × 297mm)
        - 页边距：上3cm、下3cm、左2.6cm、右2.6cm
        """
        # 获取文档的第一个（也是唯一一个）节（section）
        section = self.doc.sections[0]
        # 设置页面宽度为21厘米（A4纸宽度）
        section.page_width = Cm(21)
        # 设置页面高度为29.7厘米（A4纸高度）
        section.page_height = Cm(29.7)
        # 设置上边距为3厘米
        section.top_margin = Cm(3)
        # 设置下边距为3厘米
        section.bottom_margin = Cm(3)
        # 设置左边距为2.6厘米
        section.left_margin = Cm(2.6)
        # 设置右边距为2.6厘米
        section.right_margin = Cm(2.6)
    
    # 设置文档默认样式（私有方法）
    def _setup_styles(self):
        """
        设置文档默认样式（私有方法）
        
        设置Normal（正文）样式：
        - 字体：仿宋
        - 字号：12磅（相当于小四号字）
        - 行间距：1.5倍
        """
        # 获取Normal样式（默认正文样式）
        style = self.doc.styles['Normal']
        # 设置字体名称为仿宋
        style.font.name = self.FONT_FANGSONG
        # 设置中文字体（这一步很重要，因为Word对中文字体需要特殊处理）
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_FANGSONG)
        # 设置字体大小为12磅（小四号字）
        style.font.size = Pt(12)
        # 设置行间距为1.5倍
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # 设置文本的字体格式（私有方法）
    def _set_run_font(self, run, font_name: str, font_size: int, bold: bool = False, color: str = '000000'):
        """
        设置文本的字体格式（私有方法）
        
        参数说明：
        - run: Word中的文本运行对象（一段文字的一部分）
        - font_name: 字体名称
        - font_size: 字体大小（磅值）
        - bold: 是否加粗（True=加粗，False=不加粗）
        - color: 字体颜色（十六进制，如'FF0000'表示红色）
        """
        # 设置字体名称
        run.font.name = font_name
        # 设置中文字体（处理中文字体的特殊要求）
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        # 设置字体大小
        run.font.size = Pt(font_size)
        # 设置是否加粗
        run.font.bold = bold
        # 如果颜色不是黑色（'000000'），则设置颜色
        if color != '000000':
            # 将十六进制颜色字符串转换为RGB颜色对象并设置
            run.font.color.rgb = RGBColor.from_string(color)

    # 将阿拉伯数字日期转换为中文日期格式（私有方法）
    def _convert_date_to_chinese(self, date_str: str) -> str:
        """
        将阿拉伯数字日期转换为中文日期格式（私有方法）
        
        支持输入格式：
        - "2026年3月13日"
        - "二〇二六年三月十三日"（已为中文格式则原样返回）
        
        输出格式：
        - "二〇二六年三月十三日"
        
        参数说明：
        - date_str: 输入的日期字符串
        
        返回值：
        - 转换后的中文日期字符串
        """
        # 定义中文数字映射表
        chinese_nums = ['〇', '一', '二', '三', '四', '五', '六', '七', '八', '九']
        
        # 检查是否已经是中文格式（包含〇、一、二等字符）
        has_chinese_num = any(c in date_str for c in ['〇', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十'])
        if has_chinese_num:
            return date_str
        
        # 尝试解析日期字符串
        # 使用正则表达式提取年、月、日
        import re
        pattern = r'(\d{4})年(\d{1,2})月(\d{1,2})日'
        match = re.match(pattern, date_str)
        
        if not match:
            # 如果匹配失败，尝试其他格式或返回原字符串
            return date_str
        
        year = match.group(1)
        month = int(match.group(2))
        day = int(match.group(3))
        
        # 转换年份
        chinese_year = ''.join([chinese_nums[int(d)] for d in year])
        
        # 转换月份
        def num_to_chinese_month(num):
            if num == 1:
                return '一'
            elif num == 2:
                return '二'
            elif num == 3:
                return '三'
            elif num == 4:
                return '四'
            elif num == 5:
                return '五'
            elif num == 6:
                return '六'
            elif num == 7:
                return '七'
            elif num == 8:
                return '八'
            elif num == 9:
                return '九'
            elif num == 10:
                return '十'
            elif num == 11:
                return '十一'
            elif num == 12:
                return '十二'
            else:
                return str(num)
        
        chinese_month = num_to_chinese_month(month)
        
        # 转换日期
        def num_to_chinese_day(num):
            if num == 0:
                return '〇'
            elif num < 10:
                return chinese_nums[num]
            elif num == 10:
                return '十'
            elif num < 20:
                return '十' + chinese_nums[num % 10]
            elif num == 20:
                return '二十'
            elif num < 30:
                return '二十' + chinese_nums[num % 10]
            elif num == 30:
                return '三十'
            elif num == 31:
                return '三十一'
            else:
                return str(num)
        
        chinese_day = num_to_chinese_day(day)
        
        # 组合成最终的中文日期
        return f'{chinese_year}年{chinese_month}月{chinese_day}日'
    
    # 添加一个带格式的段落（私有方法）
    def _add_paragraph_with_font(self, text: str, font_name: str, font_size: int, 
                                   alignment=WD_ALIGN_PARAGRAPH.LEFT, bold: bool = False,
                                   first_line_indent: int = 0, color: str = '000000',
                                   space_before: int = 0, space_after: int = 0):
        """
        添加一个带格式的段落（私有方法）
        
        这是一个通用方法，用于创建各种格式的段落
        
        参数说明：
        - text: 段落的文字内容
        - font_name: 字体名称
        - font_size: 字体大小（磅值）
        - alignment: 对齐方式（左对齐、居中、右对齐）
        - bold: 是否加粗
        - first_line_indent: 首行缩进（磅值，32磅约等于2个字符）
        - color: 字体颜色
        - space_before: 段前间距（磅值）
        - space_after: 段后间距（磅值）
        
        返回值：
        - 创建好的段落对象
        """
        # 在文档中添加一个新段落
        p = self.doc.add_paragraph()
        # 设置段落的对齐方式
        p.alignment = alignment
        
        # 如果设置了首行缩进，则应用首行缩进
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Pt(first_line_indent)
        
        # 如果设置了段前间距，则应用段前间距
        if space_before > 0:
            p.paragraph_format.space_before = Pt(space_before)
        
        # 如果设置了段后间距，则应用段后间距
        if space_after > 0:
            p.paragraph_format.space_after = Pt(space_after)
        
        # 如果有文字内容，则添加文字并设置格式
        if text:
            # 在段落中添加文字（run是Word中一段文字的一部分）
            run = p.add_run(text)
            # 调用前面的方法设置字体格式
            self._set_run_font(run, font_name, font_size, bold, color)
        
        # 返回创建好的段落对象
        return p
    
   
    # 添加集团名称  汇川技术
    def add_group(self, group_name: str = '汇川技术'):
        """
        添加集团名称
        
        参数说明：
        - group_name: 集团名称，默认"汇川技术"
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置居中对齐
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 添加集团名称文字
        run = p.add_run(group_name)
        # 设置字体为宋体，18磅（小二），加粗，黑色
        self._set_run_font(run, self.FONT_SONGTI, 18, True)
        
        # 返回段落对象
        return p
    
    # 添加发文机关标志   XXX部(部门全称)文件
    def add_issuer_mark(self, issuer: str, is_red: bool = False):
        """
        添加发文机关标志
        
        参数说明：
        - issuer: 发文机关名称，如"XXX部"
        - is_red: 是否用红色（True=红色，False=黑色），默认黑色
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置居中对齐
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(10)
        # 设置单倍行距
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # 如果发文机关名称中没有"文件"二字，则自动加上
        if '文件' not in issuer:
            issuer += '文件'
        
        # 添加文字
        run = p.add_run(issuer)
        # 根据参数决定颜色，红色是'FF0000'，黑色是'000000'
        color = 'FF0000' if is_red else '000000'
        # 设置字体为宋体，42磅（初号），加粗
        self._set_run_font(run, self.FONT_SONGTI, 42, True, color)
        
        # 返回段落对象
        return p
    
    # 添加发文字号和签发人（合并版）
    def add_document_header(self, doc_number: str, signer_name: Optional[str] = None):
        """
        添加发文字号和签发人（合并版）
        
        功能说明：
        - 如果有签发人，则发文字号左对齐，签发人右对齐，在同一行
        - 如果没有签发人，则发文字号居中对齐
        
        参数说明：
        - doc_number: 发文字号，如"全球行业管理中心〔2026〕1号"
        - signer_name: 签发人姓名，可选参数，如不提供则只显示发文字号
        """
        # 发文字号起始位置偏移量（厘米），如需调整请修改此值
        DOC_NUMBER_TAB_POS = 0.2
        
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置段前间距8磅（0.5行）
        p.paragraph_format.space_before = Pt(8)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行间距为1倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        if signer_name:
            # 有签发人的情况：设置两个制表位
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            tab_stops = p.paragraph_format.tab_stops
            # 第一个制表位：控制发文字号起始位置（左对齐）
            if DOC_NUMBER_TAB_POS > 0:
                tab_stops.add_tab_stop(Cm(DOC_NUMBER_TAB_POS), WD_TAB_ALIGNMENT.LEFT)
            # 第二个制表位：签发人右对齐位置（固定15.5厘米）
            tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
            
            # 添加制表符（到第一个制表位）
            if DOC_NUMBER_TAB_POS > 0:
                run0 = p.add_run('\t')
                self._set_run_font(run0, self.FONT_FANGSONG, 16, False)
            
            # 添加发文字号
            run1 = p.add_run(doc_number)
            self._set_run_font(run1, self.FONT_FANGSONG, 16, True)
            
            # 添加制表符（到第二个制表位）
            run2 = p.add_run('\t')
            self._set_run_font(run2, self.FONT_FANGSONG, 16, False)
            
            # 添加"签发人："
            run3 = p.add_run('签发人：')
            self._set_run_font(run3, self.FONT_FANGSONG, 16, True)
            
            # 处理签发人姓名：两个字中间加空格
            processed_signer = signer_name
            if len(signer_name) == 2:
                processed_signer = f'{signer_name[0]} {signer_name[1]}'
            
            # 添加签发人姓名
            run4 = p.add_run(processed_signer)
            self._set_run_font(run4, self.FONT_KAITI, 16, True)
        else:
            # 没有签发人的情况：发文字号居中
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(doc_number)
            self._set_run_font(run, self.FONT_FANGSONG, 16, True)
        
        return p
    
    # 添加发文字号  （已废弃，请使用 add_document_header）
    def add_document_number(self, doc_number: str):
        """
        添加发文字号（已废弃，请使用 add_document_header）
        
        参数说明：
        - doc_number: 发文字号，如"全球行业管理中心〔2026〕1号"
        """
        import warnings
        warnings.warn('add_document_number 已废弃，请使用 add_document_header', DeprecationWarning, stacklevel=2)
        return self.add_document_header(doc_number, None)
    
    # 添加签发人 （已废弃，请使用 add_document_header）
    def add_signer(self, signer_name: str):
        """
        添加签发人（已废弃，请使用 add_document_header）
        
        参数说明：
        - signer_name: 签发人姓名
        """
        import warnings
        warnings.warn('add_signer 已废弃，请使用 add_document_header', DeprecationWarning, stacklevel=2)
        raise RuntimeError('add_signer 已废弃，无法单独使用，请使用 add_document_header 同时传入发文字号和签发人')
    
    # 添加黑色分隔线（版头分隔线）
    def add_black_separator(self):
        """
        添加黑色分隔线（版头分隔线）
        
        在发文机关标志和发文字号下方添加一条黑色的横线
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置居中对齐
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行间距为1倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # 添加37个"━"字符来模拟分隔线
        run = p.add_run('━' * 37)
        # 设置字体为宋体，12磅，黑色
        self._set_run_font(run, self.FONT_SONGTI, 12, False, '000000')
        
        # 返回段落对象
        return p

    # 添加密级
    def add_doc_classification(self, classification: str):
        """
        添加密级
        
        参数说明：
        - classification: 密级（机密/秘密/内部公开/外部公开）
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置左对齐
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 添加【密级：
        run1 = p.add_run('【密级：')
        # 设置字体为仿宋，16磅（3号字），黑色，加粗
        self._set_run_font(run1, self.FONT_FANGSONG, 16, True)
        
        # 添加密级内容（红色）
        run2 = p.add_run(classification)
        # 设置字体为仿宋，16磅（3号字），红色，加粗
        self._set_run_font(run2, self.FONT_FANGSONG, 16, True, 'FF0000')
        
        # 添加】
        run3 = p.add_run('】')
        # 设置字体为仿宋，16磅（3号字），黑色，加粗
        self._set_run_font(run3, self.FONT_FANGSONG, 16, True)
        
        # 返回段落对象
        return p
   
    # 添加公文标题
    def add_title(self, title: str):
        """
        添加公文标题
        
        参数说明：
        - title: 公文标题，支持用 * 手动换行，如"关于成立*ATH事业群的通知"
        """
        # 处理标题换行
        title_lines = []
        
        # 优先检测用户是否用 * 手动换行了
        if '*' in title:
            title_lines = title.split('*')
        else:
            # 没有 *，按16个字符自动换行
            max_chars_per_line = 16
            current_line = ''
            for char in title:
                if len(current_line) >= max_chars_per_line:
                    title_lines.append(current_line)
                    current_line = char
                else:
                    current_line += char
            if current_line:
                title_lines.append(current_line)
        
        # 如果只有一行，保持原来的格式
        if len(title_lines) == 1:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            run = p.add_run(title_lines[0])
            self._set_run_font(run, self.FONT_SONGTI, 22, True)
            return p
        
        # 如果有多行，按新格式设置
        for i, line in enumerate(title_lines):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 第一行：段前8磅，段后0磅，单倍行距
            if i == 0:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # 最后一行：段前0磅，段后20磅，单倍行距
            elif i == len(title_lines) - 1:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(20)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # 中间行（如果有）：段前0磅，段后0磅，单倍行距
            else:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # 添加文字
            run = p.add_run(line)
            self._set_run_font(run, self.FONT_SONGTI, 22, True)
        
        # 返回最后一个段落对象
        return p
    
    # 添加正文段落
    def add_body_paragraph(self, text: str, first_line_indent: int = 24):
        """
        添加正文段落
        
        参数说明：
        - text: 正文文字内容
        - first_line_indent: 首行缩进（磅值），默认24磅（约2个字符）
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置两端对齐
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 如果设置了首行缩进，则应用
        if first_line_indent > 0:
            p.paragraph_format.first_line_indent = Pt(first_line_indent)
        
        # 添加文字
        run = p.add_run(text)
        # 设置字体为仿宋，12磅（小四号字）
        self._set_run_font(run, self.FONT_FANGSONG, 12, False)
        
        # 返回段落对象
        return p
    
    # 添加一级标题（格式：一、）
    def add_heading_level1(self, text: str):
        """
        添加一级标题（格式：一、）
        
        参数说明：
        - text: 标题文字，如"一、工作目标"
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置悬挂缩进1.06cm（约30磅）
        p.paragraph_format.left_indent = Cm(1.06)
        p.paragraph_format.first_line_indent = Cm(-1.06)
        # 设置段前间距0.5行（约7磅）
        p.paragraph_format.space_before = Pt(7)
        # 设置段后间距0.5行（约7磅）
        p.paragraph_format.space_after = Pt(7)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 添加文字
        run = p.add_run(text)
        # 设置字体为黑体，14磅（四号字），加粗
        self._set_run_font(run, self.FONT_HEITI, 14, True)
        
        # 返回段落对象
        return p
    
    # 添加二级标题（格式：（一））
    def add_heading_level2(self, text: str):
        """
        添加二级标题（格式：（一））
        
        参数说明：
        - text: 标题文字，如"（一）主要内容"
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置段前间距0.5行（约6磅）
        p.paragraph_format.space_before = Pt(6)
        # 设置段后间距0.5行（约6磅）
        p.paragraph_format.space_after = Pt(6)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 添加文字
        run = p.add_run(text)
        # 设置字体为仿宋，12磅（小四号字），加粗
        self._set_run_font(run, self.FONT_FANGSONG, 12, True)
        
        # 返回段落对象
        return p
    
    # 添加三级标题（格式：1.）
    def add_heading_level3(self, text: str):
        """
        添加三级标题（格式：1.）
        
        参数说明：
        - text: 标题文字，如"1. 具体要求"
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置首行缩进24磅（2个字符）
        p.paragraph_format.first_line_indent = Pt(24)
        # 设置段前间距6磅
        p.paragraph_format.space_before = Pt(6)
        # 设置段后间距6磅
        p.paragraph_format.space_after = Pt(6)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 添加文字
        run = p.add_run(text)
        # 设置字体为仿宋，12磅（小四号字），不加粗
        self._set_run_font(run, self.FONT_FANGSONG, 12, False)
        
        # 返回段落对象
        return p
    
    # 添加四级标题（格式：（1））
    def add_heading_level4(self, text: str):
        """
        添加四级标题（格式：（1））
        
        参数说明：
        - text: 标题文字，如"（1）具体要求"
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置首行缩进24磅（2个字符）
        p.paragraph_format.first_line_indent = Pt(24)
        # 设置段前间距6磅
        p.paragraph_format.space_before = Pt(6)
        # 设置段后间距6磅
        p.paragraph_format.space_after = Pt(6)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 添加文字
        run = p.add_run(text)
        # 设置字体为仿宋，12磅（小四号字），不加粗
        self._set_run_font(run, self.FONT_FANGSONG, 12, False)
        
        # 返回段落对象
        return p
    
    # 添加五级标题（格式：①）
    def add_heading_level5(self, text: str):
        """
        添加五级标题（格式：①）
        
        参数说明：
        - text: 标题文字，如"①具体要求"
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置首行缩进48磅（4个字符）
        p.paragraph_format.first_line_indent = Pt(48)
        # 设置段前间距6磅
        p.paragraph_format.space_before = Pt(6)
        # 设置段后间距6磅
        p.paragraph_format.space_after = Pt(6)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 添加文字
        run = p.add_run(text)
        # 设置字体为仿宋，12磅（小四号字），不加粗
        self._set_run_font(run, self.FONT_FANGSONG, 12, False)
        
        # 返回段落对象
        return p
    
    # 添加附件说明
    def add_attachment_note(self, attachments: List[str]):
        """
        添加附件说明
        
        参数说明：
        - attachments: 附件名称列表，如['互联网服务信息登记表']
        """
        # 如果附件列表为空，则不做任何操作
        if not attachments:
            return
        
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距：只有1个附件时为24磅，多个附件时为0磅
        if len(attachments) == 1:
            p.paragraph_format.space_after = Pt(24)
        else:
            p.paragraph_format.space_after = Pt(0)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        # 设置首行缩进24磅（2个字符）
        p.paragraph_format.first_line_indent = Pt(24)
        
        # 添加"附件："文字
        run = p.add_run('附件：')
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run, self.FONT_FANGSONG, 12, False)
        
        # 如果只有一个附件
        if len(attachments) == 1:
            # 直接在同一行添加附件名称
            run2 = p.add_run(attachments[0])
            # 设置字体为仿宋，12磅（小四号）
            self._set_run_font(run2, self.FONT_FANGSONG, 12, False)
        else:
            # 如果有多个附件，每个附件单独一行
            # 遍历附件列表，从1开始编号
            for i, att in enumerate(attachments, 1):
                # 添加新段落
                p2 = self.doc.add_paragraph()
                # 设置段前间距0磅
                p2.paragraph_format.space_before = Pt(0)
                # 设置段后间距：只有最后一行才设置为24磅
                if i == len(attachments):
                    p2.paragraph_format.space_after = Pt(24)
                else:
                    p2.paragraph_format.space_after = Pt(0)
                # 设置行间距为1.5倍
                p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                # 设置首行缩进24磅（2个字符）
                p2.paragraph_format.first_line_indent = Pt(24)
                # 添加编号和附件名称
                run2 = p2.add_run(f'{i}. {att}')
                # 设置字体为仿宋，12磅（小四号）
                self._set_run_font(run2, self.FONT_FANGSONG, 12, False)
    
    # 添加成文日期
    def add_issue_date(self, date: str):
        """
        添加成文日期
        
        参数说明：
        - date: 成文日期，如"2026年3月13日"或"二〇二六年三月十三日"
        """
        # 添加一个空段落，用于调整间距
        self.doc.add_paragraph()
        
        # 添加成文日期段落
        p = self.doc.add_paragraph()
        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行距固定24磅
        p.paragraph_format.line_spacing = Pt(15)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        
        # 设置右对齐制表位（位置15.5厘米，与印发日期保持一致）
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # 添加制表符
        run1 = p.add_run('\t')
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run1, self.FONT_FANGSONG, 12, False)
        
        # 转换日期为中文格式
        converted_date = self._convert_date_to_chinese(date)
        
        # 添加成文日期文字
        run2 = p.add_run(converted_date)
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run2, self.FONT_FANGSONG, 12, False)
    
    # 添加主送机关
    def add_main_send(self, main_send: str):
        """
        添加主送机关
        
        参数说明：
        - main_send: 主送机关名称，如"集团各部门"

        """
        # 添加版记分隔线
        self._add_separator_line()

        # 添加一个新段落
        p = self.doc.add_paragraph()

        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行距固定24磅
        # p.paragraph_format.line_spacing = Pt(24)
        # p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        # 设置行间距为1倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # 添加"主送："文字
        run1 = p.add_run('主送：')
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run1, self.FONT_FANGSONG, 12, False)
        
        # 添加主送机关名称
        run2 = p.add_run(main_send)
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run2, self.FONT_FANGSONG, 12, False)
        
        # 返回段落对象
        return p
    
    # 添加抄送机关（版记部分）
    def add_copy_send(self, copy_to: str):
        """
        添加抄送机关（版记部分）
        
        参数说明：
        - copy_to: 抄送机关名称，如"集团各部门"
        """
        # 添加版记分隔线
        self._add_separator_line()
        
        # 添加一个新段落
        p = self.doc.add_paragraph()

        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行距固定24磅
        # p.paragraph_format.line_spacing = Pt(24)
        # p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # 设置行间距为1倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 添加"抄送："文字
        run1 = p.add_run('抄送：')
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run1, self.FONT_FANGSONG, 12, False)
        
        # 添加抄送机关名称，后面加上全角句号
        run2 = p.add_run(copy_to + '。')
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run2, self.FONT_FANGSONG, 12, False)
        
        # 返回段落对象
        return p
    
    # 添加印发机关和印发日期（版记部分）
    def add_print_info(self, print_org: str, print_date: str):
        """
        添加印发机关和印发日期（版记部分）
        
        参数说明：
        - print_org: 印发机关，如"集团办公室"
        - print_date: 印发日期，如"2026年3月13日"
        """

        # 添加版记末条分隔线
        self._add_separator_line()
        
        # 添加一个新段落
        p = self.doc.add_paragraph()
        
        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行距固定24磅
        # p.paragraph_format.line_spacing = Pt(24)
        # p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # 设置行间距为1倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # 设置右对齐制表位（位置15.5厘米，在可用宽度范围内）
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # 添加印发机关文字
        run1 = p.add_run(print_org)
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run1, self.FONT_FANGSONG, 12, False)
        
        # 添加制表符
        run2 = p.add_run('\t')
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run2, self.FONT_FANGSONG, 12, False)
        
        # 添加印发日期，后面加上"印发"二字
        run3 = p.add_run(f'{print_date}印发')
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run3, self.FONT_FANGSONG, 12, False)
        
        # 添加版记末条分隔线
        self._add_separator_line()
    
    # 添加版记分隔线（私有方法）
    def _add_separator_line(self):
        """
        添加版记分隔线（私有方法）
        
        在版记部分添加横线
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置段前间距0磅
        p.paragraph_format.space_before = Pt(0)
        # 设置段后间距0磅
        p.paragraph_format.space_after = Pt(0)
        # 设置行间距为1倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # 添加55个"─"字符来模拟分隔线
        run = p.add_run('─' * 55)
        # 设置字体为宋体，8磅
        self._set_run_font(run, self.FONT_SONGTI, 8, False, '000000')
    
    # 添加结尾语
    def add_closing(self, closing: str = '特此通知/通报/公示。'):
        """
        添加结尾语
        
        参数说明：
        - closing: 结尾语，默认是"特此通知/通报/公示。"
        """
        # 添加一个新段落
        p = self.doc.add_paragraph()
        # 设置段前间距2行
        p.paragraph_format.space_before = Pt(24)
        # 设置段后间距2行
        p.paragraph_format.space_after = Pt(24)
        # 设置行间距为1.5倍
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        # 设置首行缩进24磅（2个字符）
        p.paragraph_format.first_line_indent = Pt(24)
        
        # 添加结尾语文字
        run = p.add_run(closing)
        # 设置字体为仿宋，12磅（小四号）
        self._set_run_font(run, self.FONT_FANGSONG, 12, False)
        
        # 返回段落对象
        return p
    
    # 保存文档到文件
    def save(self, filepath: str):
        """
        保存文档到文件
        
        参数说明：
        - filepath: 保存的文件路径，如"example_notice.docx"
        
        返回值：
        - 保存的文件路径
        """
        # 调用Word文档对象的save方法保存文件
        self.doc.save(filepath)
        # 返回文件路径
        return filepath


# ========== 公文类型工厂函数 ==========

# 创建通知类公文
def create_notice(content: Dict) -> OfficialDocumentGenerator:
    """
    创建通知类公文
    
    参数说明：
    - content: 包含公文内容的字典
    
    返回值：
    - 配置好的公文生成器对象
    """
    # 创建公文生成器实例
    gen = OfficialDocumentGenerator()
    
    # 添加集团名称
    gen.add_group(content.get('group', '汇川技术'))
    
    # 添加发文机关标志
    gen.add_issuer_mark(content.get('issuer', ''))
    # 添加发文字号和签发人
    gen.add_document_header(
        content.get('doc_number', ''),
        content.get('signer')
    )
    
    # 添加黑色分隔线
    gen.add_black_separator()
    
    # 添加密级
    gen.add_doc_classification(content['classification'])
    
    # 添加标题
    gen.add_title(content.get('title', ''))
    
    # 遍历正文段落列表
    for para in content.get('body', []):
        # 如果段落以"一、"、"二、"..."十一、"等中文数字开头，则作为一级标题
        if re.match(r'^[一二三四五六七八九十百]+、', para):
            gen.add_heading_level1(para)
        # 如果段落以"（一）"、"（二）"..."（十一）"等中文数字括号开头，则作为二级标题
        elif re.match(r'^（[一二三四五六七八九十百]+）', para):
            gen.add_heading_level2(para)
        # 如果段落以"1."、"2."..."100."等阿拉伯数字加点开头，则作为三级标题
        elif re.match(r'^\d+\.', para):
            gen.add_heading_level3(para)
        # 如果段落以"（1）"、"（2）"..."（100）"等阿拉伯数字括号开头，则作为四级标题
        elif re.match(r'^（\d+）', para):
            gen.add_heading_level4(para)
        # 如果段落以"①"、"②"..."⑳"等圈号数字开头，则作为五级标题
        elif re.match(r'^[①-⑳]', para):
            gen.add_heading_level5(para)
        # 否则作为普通正文段落
        else:
            gen.add_body_paragraph(para)
    
    # 添加结尾语：特此通知/通报
    gen.add_closing(content['closing'])
    
    # 如果有附件，则添加附件说明
    if content.get('attachments'):
        gen.add_attachment_note(content['attachments'])
    
    # 添加成文日期
    gen.add_issue_date(content.get('date', ''))

    # 添加主送机关
    gen.add_main_send(content['main_send'])
    
    # 如果有抄送机关，则添加抄送机关
    if content.get('copy_to'):
        gen.add_copy_send(content['copy_to'])
    
    # 添加印发机关和印发日期
    gen.add_print_info(
        content.get('print_org', ''),
        content.get('print_date', '')
    )
    
    # 返回配置好的公文生成器
    return gen



# ========== 公文类型映射字典 ==========
# 这个字典将公文类型名称映射到对应的生成函数
DOCUMENT_TYPES = {
    '通知': create_notice,      # 通知类型使用create_notice函数
}


# 生成党政机关公文（主入口函数）
def generate_document(doc_type: str, content: Dict) -> str:
    """
    生成党政机关公文（主入口函数）
    
    这是外部调用的主要函数，根据公文类型调用对应的生成函数
    
    参数说明：
    - doc_type: 公文类型（通知）
    - content: 公文内容字典
    
    返回值：
    - 生成的文件路径
    """
    import os
    
    # 检查公文类型是否支持
    if doc_type not in DOCUMENT_TYPES:
        # 如果不支持，抛出错误提示
        raise ValueError(f"不支持的公文类型: {doc_type}。支持的类型: {list(DOCUMENT_TYPES.keys())}")
    
    # 根据公文类型获取对应的生成函数
    generator = DOCUMENT_TYPES[doc_type](content)
    
    # 固定保存路径
    files_dir = r'E:\97、新一轮AI探索\WriterINO\files'
    
    # 自动创建目录（如果不存在）
    os.makedirs(files_dir, exist_ok=True)
    
    # 获取发文字号和标题
    doc_number = content.get('doc_number', '未命名')
    title = content.get('title', '公文')
    
    # 删除标题中的*（*是手动换行标记，不应该出现在文件名中）
    title_for_filename = title.replace('*', '')
    
    # 清理文件名中的非法字符
    def sanitize_filename(filename):
        # Windows文件名非法字符: \/:*?"<>|
        illegal_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
        for char in illegal_chars:
            filename = filename.replace(char, '_')
        return filename
    
    # 生成文件名：{doc_number} {title}.docx
    filename = f'{doc_number} {title_for_filename}.docx'
    filename = sanitize_filename(filename)
    
    # 拼接完整路径
    final_output_path = os.path.join(files_dir, filename)
    
    # 保存文档到指定路径
    generator.save(final_output_path)
    
    # 返回文件路径
    return final_output_path


# ========== 示例代码 ==========
# 当直接运行这个文件时，会执行下面的示例代码
if __name__ == '__main__':
    # 定义示例公文内容
    example_content = {
        'classification': '内部公开',                      # 密级（机密/秘密/内部公开/外部公开）
        'group': '阿里巴巴',                               # 集团名称
        'signer': '吴泳铭',                                  # 签发人
        'issuer': '集团总裁办公室',                              # 发文机关
        'doc_number': '阿里集团〔2025〕6号',                   # 发文字号
        'title': '关于成立Alibaba Token Hub（ATH）*事业群的通知',  # 公文标题
        'body': [                                         # 正文内容列表
            '一、总体要求',
            '当前，人工智能技术正以前所未有的速度演进，AGI时代即将全面到来。大量数字化工作将由数以百亿计的AI Agent来支撑，而这些AI Agent将由模型产生的Token支撑运行，成为人类与数字世界交互的主要载体。为抢抓这一历史性机遇，全面推进集团AI战略落地，经集团董事会研究决定，现就成立Alibaba Token Hub（ATH）事业群有关事项通知如下。各事业群、各部门要充分认识到成立ATH事业群的重要战略意义，统一思想，凝聚共识，全力支持和配合ATH事业群的各项工作。',
            '二、组织架构',
            '（一）事业群定位',
            '1. Alibaba Token Hub（ATH）事业群以平行于阿里云智能事业群、电商事业群的独立事业群形态存在，是集团AI战略的核心载体。',
            '2. ATH事业群建立以"创造Token、输送Token、应用Token"为核心目标的新组织，全面统筹集团AI研发与应用。',
            '3. 事业群由集团CEO吴泳铭直接负责，强化各AI业务战略协同，以AI重塑工作方式，保持敏捷组织。',
            '（二）下属业务单元',
            '1. 通义实验室：创造领先的多模态模型，不断追求基础模型能力上限，为集团和业界提供最领先模型。',
            '2. MaaS业务线：构建高效开放的模型服务平台和技术体系，支撑全行业AI生态。',
            '3. 千问事业部：打造最好的个人AI助手。',
            '4. 悟空事业部：打造B端AI原生工作平台，将模型能力深度融入企业工作流。',
            '5. AI创新事业部：探索各类AI创新应用，快速验证新模式、新市场。',
            '三、主要职责',
            '（一）创造Token',
            '1. 通义实验室要持续加大基础研究投入，聚焦多模态大模型研发，不断突破模型能力边界。（1）加强算法创新，提升模型理解、生成、推理等核心能力；（2）推进多模态融合，实现文本、图像、音频、视频等多模态统一理解与生成；（3）优化模型效率，降低推理成本，提升服务性能。',
            '2. 建立模型迭代机制，快速响应业务需求，持续优化模型效果。（1）建立业务反馈通道，及时收集各业务线对模型的需求和建议；（2）建立模型评估体系，定期评估模型性能，持续优化改进；（3）建立模型版本管理机制，确保模型迭代的稳定性和可靠性。',
            '（二）输送Token',
            '1. MaaS业务线要构建高效、稳定、开放的模型服务平台。① 优化服务架构，提升服务可用性和稳定性；② 开放API接口，为内部业务和外部客户提供便捷的模型调用服务；③ 建立服务监控体系，实时监控服务状态，及时处理异常问题。',
            '2. 构建完善的技术支撑体系。① 建立模型部署平台，支持模型快速部署和上线；② 建立数据管理平台，为模型训练和优化提供数据支撑；③ 建立安全保障体系，确保模型服务和数据的安全。',
            '（三）应用Token',
            '1. 千问事业部要聚焦个人用户需求，打造最好的个人AI助手。（1）深入理解用户需求，持续优化产品体验；（2）拓展应用场景，覆盖学习、工作、生活等多个领域；（3）建立用户反馈机制，持续改进产品功能。',
            '2. 悟空事业部要聚焦企业客户需求，打造B端AI原生工作平台。（1）深入理解企业工作流程，将模型能力深度融入企业工作流；（2）提供定制化服务，满足不同企业的个性化需求；（3）建立客户成功体系，帮助客户实现AI应用价值最大化。',
            '3. AI创新事业部要保持敏锐的市场洞察力，快速探索各类AI创新应用。（1）建立创新孵化机制，支持内部创新项目；（2）加强与外部合作，探索新的商业模式；（3）建立快速验证机制，及时淘汰无效项目，聚焦有前景的方向。',
            '四、组织保障',
            '（一）加强组织领导',
            '成立ATH事业群筹备工作组，由集团CEO吴泳铭任组长，相关部门负责人为成员，统筹推进事业群组建各项工作。筹备工作组下设办公室，设在集团总裁办公室，负责日常协调和推进工作。各相关部门要积极配合，确保事业群组建工作顺利推进。定期召开筹备工作会议，研究解决组建过程中的重大问题。',
            '（二）完善人员配置',
            '要高度重视ATH事业群的人才队伍建设，加大人才引进和培养力度。（1）从集团内部选拔优秀人才充实到ATH事业群；（2）加大外部人才引进力度，吸引全球顶尖AI人才加盟；（3）建立人才激励机制，为人才提供良好的发展空间和待遇。建立人才梯队，确保事业群可持续发展。',
            '五、时间安排',
            '（一）筹备阶段（2025年3月-4月）',
            '1. 制定事业群组建方案，明确组织架构、职责分工和人员配置。集团总裁办公室牵头制定总体方案，各相关部门配合制定本部门对接方案。',
            '2. 召开动员大会，统一思想认识。集团召开ATH事业群成立动员大会，对相关工作进行全面部署，各相关部门及时传达会议精神。',
            '（二）组建阶段（2025年5月-6月）',
            '1. 完成人员配置，组建事业群团队。按照组建方案，完成内部人员调配和外部人员招聘，组建完整的事业群团队。',
            '2. 建立工作机制，确保事业群正常运转。建立事业群内部管理机制、沟通协调机制、决策机制等，确保事业群各项工作有序开展。',
            '（三）运营阶段（2025年7月起）',
            '1. 全面开展业务，推进AI战略落地。ATH事业群按照既定目标和计划，全面开展各项业务，推进集团AI战略落地。',
            '2. 定期评估进展，持续优化改进。建立事业群运营评估机制，定期评估事业群运营情况，及时总结经验，持续优化改进。',
            '六、有关要求',
            '（一）提高思想认识',
            '各事业群、各部门要充分认识到成立ATH事业群的重要战略意义，深刻认识到这是集团应对AGI时代挑战、抢抓AI发展机遇的重大战略举措。各级管理人员要带头学习AI知识，带头支持ATH事业群工作，发挥示范引领作用。全体员工要积极响应集团号召，主动学习和应用AI技术，为集团AI发展贡献力量。',
            '（二）强化协同配合',
            'ATH事业群的组建和运营是一项系统性工程，需要各事业群、各部门密切配合、协同推进。集团总裁办公室要做好统筹协调，建立跨部门协同工作机制。阿里云智能事业群、电商事业群等要主动与ATH事业群对接，在技术、数据、资源等方面给予支持。各职能部门要积极配合，为ATH事业群提供必要的资源保障和政策支持，形成工作合力。',
            '（三）注重实效落地',
            '要坚持问题导向、目标导向，紧密结合集团实际开展工作，防止形式主义。ATH事业群要从业务痛点出发，选择有实际需求、能带来实际价值的场景优先突破，确保各项工作取得实实在在的成效。要建立效果评估体系，定期评估工作成效，及时总结经验，不断优化完善。要鼓励创新，宽容失败，营造勇于探索、大胆实践的良好氛围。',
        ],
        'closing': '特此通知。',                          # 结尾语
        'attachments': ['《ATH事业群组建方案》','《ATH事业群职责分工》','《ATH事业群人才计划》'],          # 附件列表
        'date': '2025年3月16日',                         # 成文日期
        'main_send': '各事业群、各部门',                          # 主送机关
        'copy_to': '集团董事局',                          # 抄送机关
        'print_org': '集团总裁办公室',                        # 印发机关
        'print_date': '2025年3月16日'                   # 印发日期
    }
    
    # 调用generate_document函数生成通知类型的公文
    # 文件会保存到：E:\97、新一轮AI探索\WriterINO\files 目录下
    # 文件名格式为：{doc_number} {title}.docx
    output = generate_document('通知', example_content)
    # 打印成功消息
    print(f"公文已生成: {output}")








